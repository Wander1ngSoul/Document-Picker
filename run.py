import os
import tempfile
import win32com.client as win32
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from tkinter import Tk
from tkinter.filedialog import askopenfilenames, asksaveasfilename
import time


class DocumentMerger:
    def __init__(self):
        self.selected_files = []
        self.data_rows = []
        self.temp_files = []
        self.root = Tk()
        self.root.withdraw()

    def select_files(self):
        files = askopenfilenames(
            filetypes=[("Word files", "*.docx *.doc")],
            initialdir=os.path.expanduser("~")
        )
        self.selected_files = list(files)
        print(f"Выбрано файлов: {len(self.selected_files)}")

    def extract_data(self):
        for file in self.selected_files:
            print(f"\nОбработка файла: {file}")
            if file.lower().endswith(".doc"):
                print("Конвертация .doc в .docx...")
                file = self.convert_doc_to_docx(file)
                if file:
                    self.temp_files.append(file)
                else:
                    continue

            try:
                doc = Document(file)
                if not doc.tables:
                    print("В документе нет таблиц!")
                    continue

                found_data = False
                for table in doc.tables:
                    data_rows_found = []
                    for row_idx, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) < 8 or not any(cells): continue
                        if cells[0].isdigit() and all(c == '' or c.isdigit() or c == ' ' for c in cells): continue
                        if '№ п/п' in cells[0] or ('Фамилия' in cells[1] if len(cells) > 1 else False): continue

                        if len(cells) > 1 and ('г.р.' in cells[1] or 'рождения' in cells[1] or
                                               any(name in cells[1] for name in
                                                   ['Казаков', 'Насиров', 'Кирилл', 'Зиганшин', 'Клипов', 'Сафиуллин',
                                                    'Тухтаев'])):
                            data_rows_found.append((row_idx, cells))

                    if data_rows_found:
                        for _, cells in data_rows_found:
                            cleaned = [' '.join(c.split()) for c in cells]
                            self.data_rows.append(cleaned)
                        found_data = True
                        break

                if not found_data:
                    print("\nДанные не найдены!")
            except Exception as e:
                print(f"Ошибка: {e}")

        print(f"\nВсего собрано строк данных: {len(self.data_rows)}")

    def cleanup_temp_files(self):
        for f in self.temp_files:
            try:
                if os.path.exists(f): os.remove(f)
            except:
                pass
        self.temp_files.clear()

    def _format_cell(self, cell, text="", font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     vertical_align=WD_ALIGN_VERTICAL.TOP):
        cell.text = text
        cell.vertical_alignment = vertical_align
        cell.margin_left = Cm(0)
        cell.margin_right = Cm(0)
        cell.margin_top = Cm(0)
        cell.margin_bottom = Cm(0)

        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.alignment = alignment
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)

        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.color.rgb = None

    def create_protocol(self):
        if not self.data_rows:
            print("Нет данных для создания протокола!")
            return False

        print(f"\nСоздание протокола с {len(self.data_rows)} строками данных")
        doc = Document()

        # Настройка страницы (А4, Альбомная)
        section = doc.sections[0]
        section.orientation = 1
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        table = self.create_table(doc)

        output_path = self.save_file()
        if output_path:
            doc.save(output_path)
            print(f"Протокол сохранен: {output_path}")
            try:
                os.startfile(os.path.dirname(output_path))
            except:
                pass
            return True
        return False

    def create_table(self, doc):
        table = doc.add_table(rows=2, cols=8)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.rows[0].height = Cm(8.5)

        headers = [
            "№ п/п",
            "Фамилия, имя, отчество. Год рождения. Военный комиссариат",
            "Диагноз, категория годности к военной службе, показатель предназначения для прохождения военной службы, решение призывной комиссии муниципального образования. Статья, пункт расписания болезней и таблицы дополнительных требований",
            "Жалобы гражданина на состояние здоровья и анамнез (другие заявления гражданина и данные на него)",
            "Данные объективного исследования, специальных исследований, диагноз (по-русски)",
            "Итоговое заключение о категории годности к военной службе, показателе предназначения для прохождения военной службы. Статья, пункт расписания болезней и таблицы дополнительных требований",
            "Решение призывной комиссии. Результаты голосования комиссии",
            "Примечания",
        ]
        for i, header in enumerate(headers):
            self._format_cell(table.rows[0].cells[i], header, font_size=14, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_align=WD_ALIGN_VERTICAL.CENTER)

        for i in range(8):
            self._format_cell(table.rows[1].cells[i], str(i + 1), font_size=10, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, vertical_align=WD_ALIGN_VERTICAL.CENTER)

        for data in self.data_rows:
            row_cells = table.add_row().cells
            for col_idx in range(8):
                if col_idx == 0:
                    self._format_cell(row_cells[0], "", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                      vertical_align=WD_ALIGN_VERTICAL.TOP)
                elif col_idx < len(data):
                    align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    self._format_cell(row_cells[col_idx], data[col_idx], font_size=14, alignment=align,
                                      vertical_align=WD_ALIGN_VERTICAL.TOP)
                else:
                    self._format_cell(row_cells[col_idx], "", font_size=14, vertical_align=WD_ALIGN_VERTICAL.TOP)

        self.set_column_widths(table)
        return table

    def set_column_widths(self, table):
        widths_cm = [1.24, 3.25, 4.25, 5.25, 5, 4, 3.25, 1.44]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths_cm):
                    cell.width = Cm(widths_cm[idx])

    def save_file(self):
        return asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            initialfile="Протокол №8 от 19.04.2025",
            title="Сохранить как"
        )

    def convert_doc_to_docx(self, doc_path):
        word = None
        try:
            norm_path = os.path.normpath(doc_path)
            if not os.path.exists(norm_path): return None

            word = win32.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(norm_path)
            time.sleep(1)

            temp_dir = tempfile.gettempdir()
            base = os.path.splitext(os.path.basename(norm_path))[0].replace(".", "_").replace(" ", "_")
            temp_docx = os.path.join(temp_dir, f"{base}_converted.docx")
            if os.path.exists(temp_docx): os.remove(temp_docx)

            doc.SaveAs(temp_docx, FileFormat=16)
            doc.Close()
            word.Quit()
            return temp_docx if os.path.exists(temp_docx) else None
        except Exception as e:
            print(f"Ошибка конвертации: {e}")
            if word:
                try:
                    word.Quit()
                except:
                    pass
            return None


if __name__ == "__main__":
    merger = DocumentMerger()
    merger.select_files()
    if merger.selected_files:
        merger.extract_data()
        if merger.data_rows:
            merger.create_protocol()
            print("\nГотово!")
        else:
            print("\nДанные не найдены!")
        merger.cleanup_temp_files()
    else:
        print("Файлы не выбраны")
    input("\nНажмите Enter для выхода...")